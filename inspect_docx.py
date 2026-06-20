import docx
import os

docx_path = "Template PBL (1).docx"
output_txt = "template_structure.txt"

def inspect_docx():
    if not os.path.exists(docx_path):
        print(f"Error: {docx_path} not found.")
        return
        
    print(f"Reading DOCX from {docx_path}...")
    doc = docx.Document(docx_path)
    
    lines = []
    lines.append(f"Total Paragraphs: {len(doc.paragraphs)}")
    lines.append(f"Total Tables: {len(doc.tables)}")
    lines.append("\n=== PARAGRAPHS STRUCTURE ===")
    
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            lines.append(f"P{i+1}: [Style: {p.style.name}] -> {p.text.strip()}")
            
    lines.append("\n=== TABLES STRUCTURE ===")
    for t_idx, t in enumerate(doc.tables):
        lines.append(f"\nTable {t_idx+1}: {len(t.rows)} rows, {len(t.columns)} cols")
        for r_idx, row in enumerate(t.rows):
            cells_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            lines.append(f"  Row {r_idx+1}: {cells_text}")
            
    with open(output_txt, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
        
    print(f"Successfully inspected docx structure to {output_txt}")

if __name__ == "__main__":
    inspect_docx()
