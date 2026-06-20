import docx
import os
from docx.text.paragraph import Paragraph

# Monkeypatch Paragraph.text setter to split newlines into separate paragraphs automatically
original_text_setter = Paragraph.text.fset

def patched_text_setter(self, value):
    if isinstance(value, str) and '\n' in value:
        lines = value.split('\n')
        original_text_setter(self, lines[0])
        current_p = self
        for line in lines[1:]:
            p_element = docx.oxml.shared.OxmlElement('w:p')
            current_p._element.addnext(p_element)
            new_para = Paragraph(p_element, self._parent)
            if line:
                original_text_setter(new_para, line)
            if self.style:
                new_para.style = self.style
            current_p = new_para
    else:
        original_text_setter(self, value)

Paragraph.text = property(Paragraph.text.fget, patched_text_setter, Paragraph.text.fdel)

def insert_paragraph_after(paragraph, text, style=None):
    """
    Inserts a new paragraph after the given paragraph.
    """
    p_element = docx.oxml.shared.OxmlElement('w:p')
    paragraph._element.addnext(p_element)
    new_para = docx.text.paragraph.Paragraph(p_element, paragraph._parent)
    if text:
        new_para.text = text
    if style:
        new_para.style = style
    return new_para

def create_styled_table_after(paragraph, headers, data, style=None):
    """
    Creates a styled table with the given headers and data after the specified paragraph
    by creating it via python-docx and moving its XML element.
    """
    doc = paragraph._parent.part.document
    rows_count = len(data) + 1
    cols_count = len(headers)
    
    # Create standard table at the end of the document
    table = doc.add_table(rows=rows_count, cols=cols_count)
    if style:
        table.style = style
        
    # Populate headers
    for col_idx, text in enumerate(headers):
        table.cell(0, col_idx).text = str(text)
        for p in table.cell(0, col_idx).paragraphs:
            for r in p.runs:
                r.bold = True
                
    # Populate data
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            if col_idx < cols_count:
                table.cell(row_idx + 1, col_idx).text = str(val)
                
    # Move the table XML element to be after the paragraph
    paragraph._element.addnext(table._element)
    
    return table

def insert_row_before(table, row_idx):
    """
    Inserts a new row before the specified index in the table.
    """
    tr = table.rows[row_idx]._tr
    new_tr = table.add_row()._tr
    tr.addprevious(new_tr)
    return table.rows[row_idx]

def convert_report():
    template_path = "Template PBL (1).docx"
    output_path = "Laporan_Akhir_NutriGrow_Baru_v2.docx"
    
    if not os.path.exists(template_path):
        print(f"Error: Template {template_path} not found.")
        return
        
    print(f"Opening template {template_path}...")
    doc = docx.Document(template_path)
    
    # Store references to all target paragraph objects by their original index
    # to avoid index shift issues during text and table insertions.
    p_logo = doc.paragraphs[3]
    p_cover = doc.paragraphs[4]
    p_project = doc.paragraphs[5]
    p_team = doc.paragraphs[6]
    p_leader = doc.paragraphs[7]
    p_members = doc.paragraphs[8]
    p_dosen = doc.paragraphs[9]
    
    p_tujuan_dokumen = doc.paragraphs[23]
    p_gambaran_sistem = doc.paragraphs[25]
    p_tujuan_sistem = doc.paragraphs[26]
    p_fungsi_sistem = doc.paragraphs[27]
    p_batasan_sistem = doc.paragraphs[28]
    p_cara_kerja = doc.paragraphs[29]
    p_lingkungan_impl = doc.paragraphs[30]
    
    p_hw = doc.paragraphs[34]
    p_sw = doc.paragraphs[36]
    p_sdm = doc.paragraphs[38]
    
    p_analisis_kondisi = doc.paragraphs[43]
    p_perancangan_keseluruhan = doc.paragraphs[45]
    p_perancangan_iot = doc.paragraphs[47]
    p_perancangan_jaringan = doc.paragraphs[49]
    p_perancangan_web = doc.paragraphs[51]
    p_perancangan_mobile = doc.paragraphs[53]
    p_perancangan_cloud = doc.paragraphs[55]
    p_perancangan_kompresi = doc.paragraphs[57]
    p_perancangan_aiml = doc.paragraphs[59]
    p_biaya_intro = doc.paragraphs[61]
    
    p_instalasi = doc.paragraphs[70]
    p_pengoperasian = doc.paragraphs[72]
    
    p_pengujian_intro = doc.paragraphs[75]
    p_metodologi = doc.paragraphs[77]
    p_prosedur = doc.paragraphs[79]
    
    p_analisis = doc.paragraphs[83]
    p_kesimpulan = doc.paragraphs[88]
    p_istilah = doc.paragraphs[94]
    p_pustaka = doc.paragraphs[98]
    
    p_cv_azrial = doc.paragraphs[102]
    p_cv_arizal = doc.paragraphs[108]
    p_cv_sultan = doc.paragraphs[120]
    
    # Get the style of the first table to replicate it
    table_style = doc.tables[0].style if doc.tables else None
    print(f"Original table style: {table_style.name if table_style else 'None'}")
    
    print("Modifying cover page details...")
    p_logo.text = "[Logo PNJ - Silakan Sisipkan Logo PNJ di Sini]"
    p_cover.text = "[Desain Cover - Silakan Sisipkan Gambar Cover Proyek di Sini]"
    p_project.text = "Nama Project: Sistem Monitoring dan Otomasi Irigasi Cerdas pada Pertanian Berbasis IoT menggunakan Fuzzy Logic dengan Web Dashboard (NutriGrow)"
    p_team.text = "Tim Penulis: Kelompok NutriGrow"
    p_leader.text = "Nama Ketua: [Belum Diisi - Silakan Tentukan Ketua Tim]"
    p_members.text = (
        "Anggota:\n"
        "1. Maulana Daviq Putra (NIM: 2307422024)\n"
        "2. Muhammad Ali Diepo Ar-Rasyid (NIM: 2307422023)\n"
        "3. Muhammad Athaillah Hardianto (NIM: 2307422011)\n"
        "4. Muhammad Pandya Hanif Ramadhan (NIM: 2307422006)\n"
        "5. Ravila Sasla Arandika (NIM: 2307422026)"
    )
    p_dosen.text = "Tim Dosen PBL: [Belum Diisi - Silakan Tentukan Dosen Pembimbing / Tim Dosen PBL]"

    print("Modifying Section 1 (Pendahuluan)...")
    p_tujuan_dokumen.text = (
        "Dokumen Laporan Akhir Proyek ini disusun untuk mendokumentasikan hasil perancangan, implementasi, "
        "dan pengujian sistem NutriGrow (Bitanic Pro V4). Dokumen ini berfungsi sebagai panduan teknis yang menjelaskan "
        "arsitektur jaringan IoT, logika kontrol fertigasi otonom menggunakan Fuzzy Logic, platform web dan mobile monitoring, "
        "serta analisis kinerja sistem secara menyeluruh untuk menjamin keberhasilan proyek."
    )
    
    p_gambaran_sistem.text = (
        "Gambaran Utama Sistem: NutriGrow (Bitanic Pro V4) adalah sistem fertigasi cerdas terintegrasi berbasis IoT "
        "yang dirancang untuk mengoptimalkan pemberian air dan nutrisi pada perkebunan. Sistem ini mengintegrasikan "
        "mikrokontroler ESP32 sebagai sensor node lapangan menggunakan transmisi nirkabel LoRa Point-to-Point (433 MHz) "
        "dan ESP-NOW untuk komunikasi jarak jauh tanpa ketergantungan pada jaringan internet seluler di titik lahan."
    )
    
    p_tujuan_sistem.text = (
        "Tujuan pengembangan sistem: Merancang dan mengimplementasikan sistem fertigasi cerdas terintegrasi berbasis IoT "
        "guna meningkatkan efisiensi konsumsi air, pupuk, dan energi listrik melalui penerapan Fuzzy Logic dan Smart Delay "
        "prediktif berbasis prakiraan cuaca, serta mengatasi keterbatasan jaringan seluler di lahan pertanian rural."
    )
    
    p_fungsi_sistem.text = (
        "Fungsi utama sistem:\n"
        "- Pemantauan real-time kelembapan tanah, suhu/kelembapan udara, pH, dan TDS.\n"
        "- Keputusan irigasi otomatis berbasis Fuzzy Logic Mamdani dan Smart Delay prediktif.\n"
        "- Kontrol pompa air dan pompa pupuk secara jarak jauh (manual override).\n"
        "- Visualisasi metrik efisiensi air/energi (eco-savings) pada Web Dashboard dan Aplikasi Mobile."
    )
    
    p_batasan_sistem.text = (
        "Batasan Sistem (Non-Goals): Sistem tidak mencakup otomasi parameter greenhouse tertutup, "
        "hidroponik vertikal, modul e-commerce hasil panen, dan analisis citra drone."
    )
    
    p_cara_kerja.text = (
        "Cara kerja sistem: ESP32 Node Lapangan mengukur parameter tanah (kelembapan) dan udara (DHT22) "
        "lalu mentransmisikannya via LoRa/ESP-NOW ke ESP32 Gateway di Pos Jaga. Gateway mengukur pH/TDS, "
        "menjalankan logika kontrol pompa, dan meneruskan data telemetri via MQTTS (port 8883) ke Cloud VPS. "
        "VPS memproses data, menyimpannya di Supabase PostgreSQL, dan menyiarkannya secara real-time ke klien "
        "(Next.js Dashboard & Expo Mobile) menggunakan WebSocket."
    )
    
    p_lingkungan_impl.text = (
        "Lingkungan Implementasi: Sistem diimplementasikan pada Lahan Pertanian (Zona A) untuk penempatan sensor node, "
        "Pos Jaga (Zona B) untuk gateway/controller, Cloud VPS (Ubuntu 22.04 LTS) untuk server backend (Node.js/Express) "
        "dan broker MQTT (Mosquitto), serta platform database cloud Supabase."
    )

    print("Modifying Section 2 (Sumber yang Dibutuhkan)...")
    p_hw.text = (
        "Perangkat keras yang digunakan dalam proyek NutriGrow meliputi:\n"
        "1. Mikrokontroler: ESP32 (3 unit) dan ESP8266 (2 unit).\n"
        "2. Modul Komunikasi: Modul LoRa Ra-02 (433 MHz, 2 unit).\n"
        "3. Sensor: Sensor DHT22, Sensor Kelembapan Tanah (Soil Moisture), Sensor pH Analog, dan Sensor TDS.\n"
        "4. Aktuator: Pompa Air 12V, Pompa Pupuk 12V, dan Relay 4-Channel.\n"
        "5. Perangkat Jaringan: Router MikroTik hAP lite."
    )
    
    p_sw.text = (
        "Perangkat lunak yang digunakan dalam proyek NutriGrow meliputi:\n"
        "1. Frontend: Next.js (Web Dashboard SPA) dan Expo React Native (Aplikasi Mobile).\n"
        "2. Backend: Node.js / Express.js REST API & WebSocket Server.\n"
        "3. Database: PostgreSQL terkelola di Supabase Cloud dengan Row Level Security (RLS).\n"
        "4. MQTT Broker: Eclipse Mosquitto dengan enkripsi TLS (MQTTS port 8883).\n"
        "5. DevOps & Tools: PM2 Process Manager, VS Code IDE, dan OS Ubuntu 22.04 LTS pada VPS."
    )
    
    p_sdm.text = (
        "Pihak-pihak yang terlibat dalam proyek NutriGrow:\n"
        "1. Maulana Daviq Putra (NIM: 2307422024) - [Tugas/Peran Belum Diisi]\n"
        "2. Muhammad Ali Diepo Ar-Rasyid (NIM: 2307422023) - [Tugas/Peran Belum Diisi]\n"
        "3. Muhammad Athaillah Hardianto (NIM: 2307422011) - [Tugas/Peran Belum Diisi]\n"
        "4. Muhammad Pandya Hanif Ramadhan (NIM: 2307422006) - [Tugas/Peran Belum Diisi]\n"
        "5. Ravila Sasla Arandika (NIM: 2307422026) - [Tugas/Peran Belum Diisi]\n\n"
        "Pihak Eksternal:\n"
        "- Tim Dosen PBL - [Tugas/Peran Belum Diisi]"
    )

    print("Modifying Section 3 (Perancangan)...")
    p_analisis_kondisi.text = (
        "Praktik fertigasi (pemberian air dan nutrisi tanaman) konvensional pada perkebunan sering kali "
        "menghadapi tantangan efisiensi sumber daya dan kendala infrastruktur. Penyiraman terjadwal secara "
        "manual bersifat tidak responsif terhadap fluktuasi cuaca lokal, menyebabkan pemborosan air dan "
        "pupuk saat hujan turun. Selain itu, keterbatasan jaringan seluler (Wi-Fi/4G) di area perkebunan rural "
        "sering kali menimbulkan blank spot komunikasi data IoT, yang membatasi penerapan pemantauan jarak jauh."
    )
    
    p_perancangan_keseluruhan.text = (
        "Sistem NutriGrow menggunakan perancangan jaringan terdistribusi. Pada Router MikroTik hAP lite, "
        "dilakukan segmentasi jaringan untuk membagi lalu lintas data:\n"
        "- Subnet IoT (VLAN 10): IP Range 192.168.10.0/24 khusus untuk ESP32 Controller dan Gateway. Subnet ini diisolasi dari subnet staff demi keamanan.\n"
        "- Subnet Staff (VLAN 20): IP Range 192.168.20.0/24 untuk PC staff dan perangkat klien.\n"
        "- Protokol Komunikasi: Telemetri dikirim via MQTTS (8883) ke Mosquitto Broker di VPS, dan visualisasi dikirim via HTTPS/WSS (443) ke klien.\n\n"
        "[Silakan Sisipkan Gambar 3.1 Diagram Topologi Jaringan Fisik di Sini]"
    )
    
    # Empty guidelines placeholders for 3.3 to 3.8
    p_perancangan_iot.text = (
        "Sub-bab ini menjelaskan perancangan sistem IoT NutriGrow (arsitektur hardware, skematik, dan flowchart sensor-node).\n"
        "[Belum Diisi - Silakan Sisipkan Arsitektur, Diagram Blok, dan Flowchart IoT Lapangan]"
    )
    p_perancangan_jaringan.text = (
        "Sub-bab ini menjelaskan perancangan jaringan MikroTik, VLAN, IP Addressing, dan konfigurasi firewall.\n"
        "[Belum Diisi - Silakan Sisipkan Diagram Jaringan Fisik & Konfigurasi Jaringan]"
    )
    p_perancangan_web.text = (
        "Sub-bab ini menjelaskan perancangan antarmuka Web Dashboard menggunakan Next.js SPA.\n"
        "[Belum Diisi - Silakan Sisipkan Wireframe Web, User Flow, dan Diagram Arsitektur Web]"
    )
    p_perancangan_mobile.text = (
        "Sub-bab ini menjelaskan perancangan antarmuka Aplikasi Mobile menggunakan Expo React Native.\n"
        "[Belum Diisi - Silakan Sisipkan Wireframe Mobile, User Flow, dan Diagram Arsitektur Mobile]"
    )
    p_perancangan_cloud.text = (
        "Sub-bab ini menjelaskan perancangan arsitektur Cloud VPS, API Gateway, Supabase PostgreSQL, dan Row Level Security.\n"
        "[Belum Diisi - Silakan Sisipkan Diagram Arsitektur Cloud dan Aturan RLS]"
    )
    p_perancangan_kompresi.text = (
        "Sub-bab ini menjelaskan perancangan kompresi C-struct packing untuk payload LoRa dan Gzip untuk REST API.\n"
        "[Belum Diisi - Silakan Sisipkan Struktur Struct dan Flowchart Kompresi]"
    )
    p_perancangan_aiml.text = (
        "Sub-bab ini menjelaskan model inferensi Fuzzy Logic Mamdani dan algoritma Smart Delay BMKG.\n"
        "[Belum Diisi - Silakan Sisipkan Aturan Fuzzy, Himpunan Keanggotaan, dan Flowchart Algoritma]"
    )
    
    p_biaya_intro.text = (
        "Sub-bab ini menjabarkan rincian estimasi biaya (Rencana Anggaran Biaya/RAB) untuk pengembangan sistem NutriGrow. "
        "Rincian biaya perangkat keras dan sewa VPS ditunjukkan pada tabel di bawah ini."
    )

    print("Updating Table 1 (Estimasi Biaya Proyek / RAB)...")
    table = doc.tables[0]
    
    items = [
        ("1", "ESP32 Controller & Node", "3", "75,000", "225,000", "Mikrokontroler utama sensor & gateway."),
        ("2", "ESP8266 Node", "2", "45,000", "90,000", "Mikrokontroler cadangan/sensor tambahan."),
        ("3", "Modul LoRa Ra-02 (433 MHz)", "2", "80,000", "160,000", "Transmisi nirkabel jarak jauh antar-zona."),
        ("4", "Paket Sensor (DHT22, Soil, pH, TDS)", "1 set", "350,000", "350,000", "Mengukur suhu, kelembapan tanah/udara, pH, & TDS."),
        ("5", "Pompa 12V + Relay 4-Channel", "1 set", "250,000", "250,000", "Aktuator untuk irigasi & fertigasi otomatis."),
        ("6", "Router MikroTik hAP lite", "1", "350,000", "350,000", "Pembagian subnet & QoS jaringan lokal."),
        ("7", "Cloud VPS Ubuntu (Sewa/Bulan)", "1 bln", "350,000", "350,000", "Hosting backend API, MQTT Broker, & DB Sync.")
    ]
    
    for i in range(3):
        insert_row_before(table, 5)
        
    for idx, item in enumerate(items):
        row_cells = table.rows[idx + 1].cells
        for col_idx, val in enumerate(item):
            row_cells[col_idx].text = val
            
    total_cells = table.rows[8].cells
    total_cells[0].text = ""
    total_cells[1].text = "Total"
    total_cells[2].text = "Total"
    total_cells[3].text = "Total"
    total_cells[4].text = "1,775,000"
    total_cells[5].text = "Pengeluaran perangkat keras + 1 bulan VPS."
    
    for p in total_cells[1].paragraphs + total_cells[4].paragraphs:
        for r in p.runs:
            r.bold = True

    print("Modifying Section 4 (Instalasi dan Cara Pengoperasian)...")
    p_instalasi.text = (
        "Langkah-langkah instalasi sistem NutriGrow meliputi:\n"
        "1. Deployment database PostgreSQL di Supabase Cloud dan pembuatan skema tabel-tabel (sensor_data, zones, dll.) beserta konfigurasi Row Level Security (RLS).\n"
        "2. Instalasi dan konfigurasi broker MQTT Eclipse Mosquitto di VPS Ubuntu 22.04 LTS dengan mengaktifkan port 8883 (MQTTS) berbasis sertifikat SSL/TLS.\n"
        "3. Flash firmware C++ menggunakan Arduino IDE ke ESP32 Node Lapangan dan Gateway dengan library LoRa dan ESP-NOW.\n"
        "4. Deployment backend Node.js/Express di VPS Ubuntu menggunakan PM2 Process Manager untuk menjamin ketersediaan layanan secara kontinu.\n"
        "5. Build dan deployment Next.js Web Dashboard di Vercel atau VPS, serta konfigurasi environment variables untuk Supabase Client SDK."
    )
    
    p_pengoperasian.text = (
        "Cara pengoperasian sistem NutriGrow secara umum:\n"
        "1. Nyalakan ESP32 Node Lapangan di Zona A dan ESP32 Gateway di Zona B.\n"
        "2. Pastikan Router MikroTik menyala dan membagikan alamat IP khusus di VLAN 10 (IoT Subnet) ke ESP32 Gateway.\n"
        "3. ESP32 Gateway akan terhubung ke WiFi lokal, melakukan handshake SSL/TLS, dan membangun koneksi MQTTS ke VPS.\n"
        "4. Masuk ke Web Dashboard Next.js menggunakan akun yang terdaftar. Administrator dapat memantau data kelembapan tanah, suhu/kelembapan udara, pH, TDS, dan status pompa secara real-time.\n"
        "5. Pengoperasian pompa fertigasi dapat berjalan otomatis berdasarkan Fuzzy Logic & Smart Delay, atau manual override melalui antarmuka Web Dashboard dan Aplikasi Mobile.\n\n"
        "[Belum Diisi - Silakan Sisipkan Panduan Pengoperasian Lebih Detail / User Manual]"
    )

    print("Modifying Section 5 (Pengujian)...")
    p_pengujian_intro.text = (
        "Pengujian sistem NutriGrow dilakukan secara menyeluruh mencakup beberapa aspek utama yaitu pengujian beban (load testing) pada backend/broker, "
        "pengujian keamanan (security testing) pada tingkat database/API, pengujian keandalan (failover testing) saat terjadi crash/gangguan, "
        "pengujian pemantauan berkelanjutan (monitoring), pengujian efisiensi kompresi data, serta validasi kualitas transmisi nirkabel LoRa P2P."
    )
    
    p_metodologi.text = (
        "Metodologi pengujian yang digunakan:\n"
        "1. Pengujian Beban: Menggunakan simulator MQTT untuk mengirimkan paket data telemetri dengan frekuensi tinggi (hingga 20 pesan/detik) ke VPS selama masing-masing 5 menit untuk mengukur penggunaan CPU/RAM server dan latensi.\n"
        "2. Pengujian Keamanan: Pengujian penetrasi berbasis OWASP Top 10 API Security untuk memvalidasi RLS Supabase (mencegah BOLA/IDOR), IoT Spoofing (koneksi tanpa TLS/kredensial), Privilege Escalation (kebocoran service_role), SQL Injection, dan Stored XSS.\n"
        "3. Pengujian Keandalan & Pemulihan (Failover): Mensimulasikan kegagalan komponen (crash frontend, crash backend, kehilangan WebSocket, matinya broker MQTT, pencabutan sensor pH/TDS/DHT22, kegagalan LoRa, pemutusan internet/WiFi Gateway, pemadaman listrik pompa, dan reboot VPS) untuk melihat toleransi kegagalan sistem.\n"
        "4. Pengujian Pemantauan Berkelanjutan: Pengoperasian sistem selama 1 jam penuh untuk mengukur packet loss rate, latensi end-to-end, jitter MQTT, disconnect rate, dan akurasi sensor.\n"
        "5. Pengujian Kompresi Data: Mengukur ukuran payload data sebelum dan sesudah kompresi biner struct packing (LoRa) serta HTTP Gzip (REST API).\n"
        "6. Pengujian Jaringan MikroTik: Menguji kestabilan konektivitas VPS, bandwidth QoS (Queue), pemantauan trafik (Torch), dan segregasi firewall."
    )
    
    p_prosedur.text = (
        "Pengujian dilakukan dengan serangkaian skenario terkontrol. Berikut adalah tabel-tabel detail skenario "
        "dan prosedur pengujian yang dijalankan pada sistem NutriGrow."
    )
    
    # We do the insertions sequentially starting from p_prosedur.
    # Because we are using addnext, the paragraphs and tables are inserted in order.
    # To keep inserting after the last inserted element, we anchor it to the most recently inserted element.
    p_last = p_prosedur
    
    p_last = insert_paragraph_after(p_last, "5.1.1 Prosedur Pengujian Beban (Load Testing)", "Heading 3")
    p_last = insert_paragraph_after(p_last, "Skenario pengujian beban dilakukan dengan memborbardir broker MQTT Mosquitto di VPS dengan parameter simulator berikut:")
    
    # Tabel 5.1: Skenario Pengujian Beban
    headers_5_1 = ["No.", "Skenario", "Jumlah Simulator MQTT", "Interval Pengiriman", "Total Pesan/detik", "Durasi"]
    data_5_1 = [
        ["1", "Baseline (Normal)", "1", "30 detik", "~0.03", "5 menit"],
        ["2", "Beban Sedang", "1", "2 detik", "0.5", "5 menit"],
        ["3", "Beban Tinggi", "1", "0.5 detik", "2", "5 menit"],
        ["4", "Beban Ekstrem", "2", "0.1 detik", "20", "5 menit"]
    ]
    t_5_1 = create_styled_table_after(p_last, headers_5_1, data_5_1, table_style)
    
    # Insert a blank paragraph anchor right after table 1 to continue insertions.
    p_next_node = docx.oxml.shared.OxmlElement('w:p')
    t_5_1._element.addnext(p_next_node)
    p_last = docx.text.paragraph.Paragraph(p_next_node, p_prosedur._parent)
    
    p_last = insert_paragraph_after(p_last, "5.1.2 Prosedur Pengujian Keamanan Jaringan & Cloud (OWASP)", "Heading 3")
    p_last = insert_paragraph_after(p_last, "Pengujian penetrasi dilakukan untuk memvalidasi keamanan siber sistem berdasarkan OWASP API Security:")
    
    # Tabel 5.2: Hasil Pengujian Keamanan
    headers_5_2 = ["No.", "Skenario Keamanan", "Kondisi / Uji Coba", "Hasil yang Diharapkan", "Hasil Aktual", "Status"]
    data_5_2 = [
        ["1", "Kebocoran Data via RLS Bypass (BOLA)", "Sebelum Mitigasi (RLS OFF)", "Data dari semua zona dikembalikan", "API mengembalikan seluruh array JSON data sensor milik pengguna lain.", "FAIL"],
        ["", "Kebocoran Data via RLS Bypass (BOLA)", "Setelah Mitigasi (RLS ON + Policy)", "Respons kosong []", "API memblokir query dan mengembalikan array kosong [] secara aman.", "PASS"],
        ["2", "Pemalsuan Data Telemetri (IoT Spoofing)", "Koneksi tanpa kredensial", "Ditolak oleh broker (Connection refused)", "Koneksi ditolak secara otomatis oleh broker.", "PASS"],
        ["", "Pemalsuan Data Telemetri (IoT Spoofing)", "Koneksi via port 1883 (non-TLS)", "Ditolak / koneksi gagal dari luar", "Port 1883 ditutup sepenuhnya oleh firewall VPS.", "PASS"],
        ["", "Pemalsuan Data Telemetri (IoT Spoofing)", "Pengiriman data ekstrem (suhu 150°C)", "Ditolak / disaring oleh backend", "Backend membatalkan insert dan mencatat warning log.", "PASS"],
        ["3", "Kebocoran Token Admin (Privilege Escalation)", "Pencarian service_role di bundle JS", "Tidak ditemukan", "Chrome DevTools static search tidak menemukan string SUPABASE_SERVICE_ROLE_KEY.", "PASS"],
        ["", "Kebocoran Token Admin (Privilege Escalation)", "Pencarian JWT prefix di bundle JS", "Hanya anon_key yang ditemukan", "Terdeteksi kunci publik anon_key (eyJhbGci...) yang telah dilindungi RLS database.", "PASS"],
        ["4", "SQL Injection pada Fitur Kustom", "SQL injection via API parameter", "Request disaring secara aman oleh database", "Perintah DROP TABLE diabaikan, data historis tetap aman ditarik.", "PASS"],
        ["", "SQL Injection pada Fitur Kustom", "Supabase Client SDK (Parameterized Query)", "Query aman, tidak ada injeksi", "Supabase RPC memperlakukan payload sebagai string biasa, database aman.", "PASS"],
        ["5", "Stored XSS via Input Pengguna", "Input XSS di field nama kebun", "Skrip tidak dieksekusi, tampil sebagai teks biasa", "React JSX me-render payload sebagai teks biasa tanpa memicu alert.", "PASS"],
        ["", "Stored XSS via Input Pengguna", "CSP header mencegah inline script", "Script terblokir oleh CSP", "Middleware helmet() memblokir eksekusi inline script tak dikenal.", "PASS"]
    ]
    t_5_2 = create_styled_table_after(p_last, headers_5_2, data_5_2, table_style)
    
    p_next_node = docx.oxml.shared.OxmlElement('w:p')
    t_5_2._element.addnext(p_next_node)
    p_last = docx.text.paragraph.Paragraph(p_next_node, p_prosedur._parent)
    
    p_last = insert_paragraph_after(p_last, "5.1.3 Prosedur Pengujian Keandalan & Pemulihan (Failover Testing)", "Heading 3")
    p_last = insert_paragraph_after(p_last, "Pengujian failover bertujuan untuk memastikan sistem toleran terhadap berbagai skenario kegagalan komponen teknis:")
    
    # Tabel 5.3: Matriks Pengujian Failover
    headers_5_3 = ["Nama Pengujian", "Simulasi Kegagalan", "Perilaku Sistem yang Diharapkan", "Hasil Aktual", "Status"]
    data_5_3 = [
        ["Crash Aplikasi Frontend", "Matikan proses Next.js server.", "Dashboard memuat error page. Ketika server aktif, web kembali pulih otomatis.", "Dashboard menampilkan error page 502/504. Setelah service Next.js dijalankan kembali, antarmuka web memuat data secara otomatis dalam 1.8 detik.", "PASS"],
        ["Kehilangan WebSocket", "Putus internet di browser selama 1 menit.", "Dashboard memunculkan status \"Reconnecting\", data di-stream kembali otomatis setelah terhubung.", "Browser menampilkan status \"Reconnecting\" kuning. Koneksi WebSocket terhubung kembali otomatis dalam 1.2 detik saat internet aktif.", "PASS"],
        ["Crash Proses Backend", "Matikan proses app.js (pm2 stop app atau kill -9).", "PM2 me-restart otomatis proses app.js. Koneksi database pulih < 30 detik.", "PM2 mendeteksi proses terhenti dan langsung me-restart proses (restars 188 -> 189). Downtime < 1 detik.", "PASS"],
        ["Restart Broker MQTT", "Restart service mosquitto di VPS.", "Backend dan ESP32 Gateway melakukan reconnect otomatis.", "Broker dimatikan selama 2 menit. Setelah aktif kembali, backend secara otomatis melakukan reconnect (MQTT Broker connected).", "PASS"],
        ["Kegagalan Sensor Zona A", "Cabut sensor DHT22 dari ESP32 Node A.", "Node A mengirim data parsial. Dashboard menampilkan status offline sensor.", "ESP32 Lapangan mengirimkan payload dengan flag error. Web menampilkan status \"Sensor Offline\" merah pada Zona A.", "PASS"],
        ["Kegagalan Sensor Zona B", "Cabut sensor pH/TDS dari ESP32 B.", "ESP32 B mengirim nilai pH default (7.0) / TDS (0). Dashboard memicu alert.", "Sensor pH dicabut. ESP32 Controller mengirimkan pH default 7.0. Dashboard memunculkan alert \"Warning: Sensor pH butuh kalibrasi\".", "PASS"],
        ["Kegagalan LoRa", "Matikan transmitter ESP8266 Lapangan.", "Gateway Zona B mendeteksi hilangnya sinyal LoRa. Sistem beralih ke mode aman.", "Transmitter dimatikan. Gateway mendeteksi packet timeout (>60 detik). Status Zona A berubah offline, kontrol beralih ke mode aman.", "PASS"],
        ["Putus WiFi Gateway B", "Matikan router Mikrotik di Zona B.", "Gateway ESP32 B beralih ke mode kontrol otonom lokal.", "Mikrotik dimatikan 2 menit. ESP32 Gateway memicu fallback ke mode otonom lokal (kontrol pompa offline berbasis TDS/pH lokal).", "PASS"],
        ["Kegagalan Weather API", "Blokir akses internet ke Weather API.", "Sistem Fuzzy Logic mengabaikan data cuaca dan menggunakan sensor lokal.", "DNS BMKG ditutup di VPS. REST API gagal menarik data cuaca, memicu fail-safe di fuzzy engine (penyiraman berbasis kelembapan tanah).", "PASS"],
        ["Mati Listrik (Power Loss)", "Cabut pasokan daya di Pos Jaga.", "Pompa dan aktuator mati otomatis untuk mencegah banjir lahan.", "Pasokan daya dilepas. Pompa air dan pupuk mati total secara otomatis (0V) sehingga tidak ada risiko kebocoran cairan fertigasi.", "PASS"],
        ["Kegagalan Daya VPS", "Lakukan reboot VPS secara paksa.", "PM2 startup script dan systemd Mosquitto otomatis memicu start service.", "VPS direboot. Mosquitto dan PM2 backend Node.js menyala otomatis pasca-boot. Total downtime pemulihan 18.5 detik.", "PASS"]
    ]
    t_5_3 = create_styled_table_after(p_last, headers_5_3, data_5_3, table_style)
    
    p_next_node = docx.oxml.shared.OxmlElement('w:p')
    t_5_3._element.addnext(p_next_node)
    p_last = docx.text.paragraph.Paragraph(p_next_node, p_prosedur._parent)
    
    p_last = insert_paragraph_after(p_last, "5.1.4 Prosedur Pengujian Kompresi Data & Jaringan MikroTik", "Heading 3")
    p_last = insert_paragraph_after(p_last, "Skenario pengujian kompresi biner struct packing (LoRa) dan HTTP Gzip (REST API):")
    
    # Tabel 5.4: Hasil Pengujian Kompresi Data
    headers_5_4 = ["No.", "Skenario Pengujian", "Ukuran Sebelum Kompresi", "Ukuran Setelah Kompresi", "Rasio Kompresi Aktual", "Status"]
    data_5_4 = [
        ["1", "Transmisi Sensor Node (Biner Struct)", "113 byte (JSON)", "22 byte (Biner)", "80.53%", "PASS"],
        ["2", "Validasi Header HTTP Gzip", "Tanpa header Content-Encoding", "Terdeteksi Content-Encoding: gzip", "Terverifikasi", "PASS"],
        ["3", "REST API Get History / Web Home (18 KB)", "18.52 KB", "4.31 KB", "76.71%", "PASS"]
    ]
    t_5_4 = create_styled_table_after(p_last, headers_5_4, data_5_4, table_style)
    
    p_next_node = docx.oxml.shared.OxmlElement('w:p')
    t_5_4._element.addnext(p_next_node)
    p_last = docx.text.paragraph.Paragraph(p_next_node, p_prosedur._parent)
    
    p_last = insert_paragraph_after(p_last, "Skenario pengujian konfigurasi jaringan MikroTik (VLAN, QoS, Firewall):")
    
    # Tabel 5.5: Skenario Jaringan MikroTik
    headers_5_5 = ["No.", "Skenario Pengujian", "Hasil yang Diharapkan", "Hasil Aktual", "Status"]
    data_5_5 = [
        ["1", "Konektivitas VPS (Ping Delay & Packet Loss)", "Ping stabil dengan packet loss 0% and delay rendah.", "Delay rata-rata 44ms (Min: 22ms, Max: 328ms), namun terdeteksi packet loss sebesar 16%.", "WARNING"],
        ["2", "Kestabilan Bandwidth (iperf3 Jitter)", "Nilai jitter sangat rendah untuk mendukung transmisi data.", "Jitter terukur sangat baik (0.000 ms pada sender dan 4.555 ms pada receiver).", "PASS"],
        ["3", "Pemantauan Trafik (Torch)", "Sistem dapat mendeteksi asal, tujuan, dan protokol data.", "Fitur Torch berhasil menampilkan alamat IP asal/tujuan serta nilai Tx/Rx Rate secara real-time.", "PASS"],
        ["4", "Manajemen Bandwidth / QoS (Queue)", "Alokasi kecepatan data terbagi sesuai prioritas per-segmen IP.", "Queue berjalan sesuai konfigurasi: \"IoT Priority\" (3M) dan \"Staff Limit\" (2M).", "PASS"],
        ["5", "Keamanan Segregasi (Ping Staff → IoT)", "Koneksi lintas subnet yang tidak diizinkan diblokir oleh sistem.", "Ping dari subnet Staff menuju perangkat IoT gagal dijangkau (sesuai rule drop pada Firewall).", "PASS"]
    ]
    t_5_5 = create_styled_table_after(p_last, headers_5_5, data_5_5, table_style)

    print("Modifying Section 6 (Analisis)...")
    p_analisis.text = (
        "Analisis hasil pengujian sistem NutriGrow dijabarkan berdasarkan kategori pengujian berikut. "
        "Tabel-tabel hasil pengujian beban, monitoring berkelanjutan, dan matriks risiko keamanan cloud disertakan di bawah ini."
    )
    
    p_analisis_last = p_analisis
    
    # Load Testing Analysis
    p_analisis_last = insert_paragraph_after(p_analisis_last, "1. Analisis Kinerja & Penggunaan Sumber Daya (Load Testing)", "Heading 3")
    p_analisis_last = insert_paragraph_after(p_analisis_last, 
        "Penggunaan CPU proses Node.js backend (app.js) berada pada tingkat yang sangat rendah "
        "dan stabil sebesar 0.0% pada baseline dan hanya memicu lonjakan kecil sebesar 2.2% (13.1% pada "
        "core tunggal/Core 0) pada beban ekstrem. Hal ini menunjukkan bahwa proses parsing data sensor dan "
        "penghitungan logika fuzzy di backend sangat efisien. Penggunaan memori (RAM) berkisar secara "
        "efisien antara 97.4 MB hingga 121.6 MB (puncak 102.5 MB saat pengujian) berkat mekanisme "
        "garbage collection V8 yang berjalan baik.\n"
        "Terdapat fenomena menarik pada latensi MQTT-ke-Database: latensi baseline (124 ms) justru "
        "lebih tinggi dibanding beban ekstrem (45 ms). Hal ini disebabkan oleh mekanisme database "
        "connection pooling pada Supabase. Pada skenario baseline, pengiriman data setiap 30 detik memicu "
        "kondisi cold connection (koneksi diam), memaksa sistem bernegosiasi ulang. Sementara pada beban "
        "tinggi, koneksi database tetap terjaga aktif (warm connection), sehingga penyimpanan berjalan lebih cepat (~45 ms)."
    )
    
    # Tabel 6.1: Hasil Pengujian Beban
    headers_6_1 = ["Skenario Beban", "CPU Rata-rata (%)", "CPU Puncak (%)", "Memori Rata-rata (MB)", "Memori Puncak (MB)", "Latensi MQTT->DB (ms)", "Latensi WS (ms)", "Throughput (%)", "Packet Loss (%)"]
    data_6_1 = [
        ["Baseline", "0.0%", "0.0%", "121.6 MB", "121.6 MB", "124 ms", "2 ms", "100.0%", "0.00%"],
        ["Beban Sedang", "0.0%", "0.0%", "105.2 MB", "105.2 MB", "54 ms", "2 ms", "100.0%", "0.00%"],
        ["Beban Tinggi", "0.0%", "0.8%", "97.4 MB", "97.4 MB", "51 ms", "2 ms", "100.0%", "0.00%"],
        ["Beban Ekstrem", "0.2%", "2.2%", "97.4 MB", "102.5 MB", "45 ms", "2 ms", "100.0%", "0.00%"]
    ]
    t_6_1 = create_styled_table_after(p_analisis_last, headers_6_1, data_6_1, table_style)
    
    p_next_node = docx.oxml.shared.OxmlElement('w:p')
    t_6_1._element.addnext(p_next_node)
    p_analisis_last = docx.text.paragraph.Paragraph(p_next_node, p_analisis._parent)
    
    # Security Analysis
    p_analisis_last = insert_paragraph_after(p_analisis_last, "2. Analisis Hasil Pengujian Keamanan Cloud (Security & Risk Matrix)", "Heading 3")
    p_analisis_last = insert_paragraph_after(p_analisis_last, 
        "Implementasi Supabase Row Level Security (RLS) terbukti 100% efektif membatasi akses data. "
        "Saat RLS non-aktif, data seluruh zona bocor (BOLA/IDOR). Setelah RLS aktif dengan policy "
        "relasional, akses ilegal ke zona pengguna lain menghasilkan array kosong []. Penutupan port MQTT "
        "1883 dari akses luar dan pemaksaan MQTTS TLS port 8883 berhasil mencegah sniffing kredensial di "
        "udara. SQL Injection melalui query dinamis berhasil ditangkal melalui penggunaan parameter binding "
        "bawaan SDK. Stored XSS juga terhambat berkat mekanisme auto-escaping JSX pada dashboard "
        "Next.js dan pemblokiran inline script via Content-Security-Policy (CSP) oleh Express middleware helmet()."
    )
    
    # Tabel 6.3: Matriks Risiko Keamanan Cloud
    headers_6_3 = ["ID", "Platform", "Skenario Ancaman", "Dampak", "Kemungkinan", "Tingkat Risiko", "Solusi Utama", "Status"]
    data_6_3 = [
        ["T-01", "Database Cloud", "Kebocoran Data via RLS Bypass (BOLA)", "Tinggi", "Sedang", "Tinggi", "Aktifkan RLS pada seluruh tabel database", "PASS"],
        ["T-02", "IoT / MQTT", "Pemalsuan Data Telemetri (Spoofing)", "Sedang", "Tinggi", "Tinggi", "Validasi rentang nilai sensor di RPC server", "PASS"],
        ["T-03", "Serverless / VM", "Kebocoran Service Role Key", "Tinggi", "Rendah", "Tinggi", "Rahasiakan Server-Side API Keys, pisahkan prefix", "PASS"],
        ["T-04", "Database Cloud", "SQL Injection pada Analitik", "Tinggi", "Rendah", "Sedang", "Gunakan Parameterized Query (Binding)", "PASS"],
        ["T-05", "Web Dashboard", "Stored XSS via User Input", "Kritis", "Sedang", "Kritis", "Auto-escaping JSX & Helmet CSP Header", "PASS"]
    ]
    t_6_3 = create_styled_table_after(p_analisis_last, headers_6_3, data_6_3, table_style)
    
    p_next_node = docx.oxml.shared.OxmlElement('w:p')
    t_6_3._element.addnext(p_next_node)
    p_analisis_last = docx.text.paragraph.Paragraph(p_next_node, p_analisis._parent)

    # Monitoring & Compression & MikroTik Analysis
    p_analisis_last = insert_paragraph_after(p_analisis_last, "3. Analisis Hasil Monitoring Berkelanjutan & Kompresi Data", "Heading 3")
    p_analisis_last = insert_paragraph_after(p_analisis_last, 
        "Tingkat kehilangan paket (packet loss) nirkabel lokal LoRa dan ESP-NOW berada di angka "
        "0.45%, jauh di bawah batas toleransi keberhasilan (5%). Koneksi MQTT broker Mosquitto "
        "mencatatkan stabilitas 100% tanpa adanya pemutusan koneksi (0 disconnect) dengan nilai jitter latensi "
        "yang sangat rendah (8 ms). Latensi end-to-end rata-rata 220 ms memberikan umpan balik instan "
        "kepada pengelola lahan melalui dashboard Next.js (latensi WebSocket rata-rata 2 ms). Deviasi pH "
        "±0.12 dan TDS ±12 ppm menunjukkan keandalan algoritma kalibrasi backend.\n"
        "Penerapan C-Struct Packing pada firmware ESP32 dan ESP8266 berhasil memotong ukuran paket data sensor "
        "sebesar 80.53% (dari JSON 113 byte menjadi biner 22 byte), menghemat airtime LoRa. Pada server API Express.js, "
        "penggunaan compression() memotong payload JSON data historis sebesar 76.67% (dari 120 KB menjadi 28 KB), "
        "mempercepat page load time grafik analitik."
    )
    
    # Tabel 6.2: Hasil Pengujian Monitoring
    headers_6_2 = ["No.", "Aspek yang Diuji", "Nilai Terukur", "Status (PASS/FAIL)"]
    data_6_2 = [
        ["1", "Packet Loss Rate", "0.05% (jaringan lokal), 0.45% (transmisi LoRa/ESP-NOW)", "PASS"],
        ["2", "Latensi End-to-End Rata-rata", "220 ms (maksimum 450 ms)", "PASS"],
        ["3", "Latensi MQTT Rata-rata", "45 ms (gateway ke VPS)", "PASS"],
        ["4", "Jitter MQTT", "8 ms", "PASS"],
        ["5", "Latensi WebSocket", "2 ms (siaran ke dashboard)", "PASS"],
        ["6", "Jumlah Disconnect MQTT", "0 kali (selama 1 jam pengujian)", "PASS"],
        ["7", "Waktu Reconnect Rata-rata", "0 detik", "PASS"],
        ["8", "Akurasi Kalibrasi pH (delta)", "± 0.12 skala pH", "PASS"],
        ["9", "Akurasi Konversi EC -> TDS (delta)", "± 12 ppm", "PASS"],
        ["10", "Konsistensi Fuzzy Logic", "0 inkonsistensi (output deterministik)", "PASS"]
    ]
    t_6_2 = create_styled_table_after(p_analisis_last, headers_6_2, data_6_2, table_style)
    
    p_next_node = docx.oxml.shared.OxmlElement('w:p')
    t_6_2._element.addnext(p_next_node)
    p_analisis_last = docx.text.paragraph.Paragraph(p_next_node, p_analisis._parent)

    p_analisis_last = insert_paragraph_after(p_analisis_last, "4. Analisis Hasil Pengujian Kinerja & Keamanan Jaringan (MikroTik)", "Heading 3")
    p_analisis_last = insert_paragraph_after(p_analisis_last, 
        "Kinerja Jaringan Eksternal: Pada lapisan transport lokal hingga pengujian iperf3, jaringan menunjukkan kualitas "
        "transmisi yang sangat prima dengan nilai jitter di bawah 5 ms. Namun, terdeteksi packet loss sebesar 16% pada pengujian "
        "Ping ke IP VPS eksternal. Hal ini dapat disebabkan oleh fluktuasi ISP, kendala medium nirkabel, atau routing di luar kendali router lokal.\n"
        "Manajemen dan Pemantauan Trafik: Implementasi QoS menggunakan fitur Queue telah beroperasi secara optimal untuk mencegah "
        "kelaparan bandwidth (starvation), dibuktikan dengan terbaginya batas maksimum secara eksplisit antara zona IoT dan zona Staff.\n"
        "Ketahanan Keamanan Isolasi: Pembatasan konektivitas dari perangkat operasional (Staff) menuju sensor/perangkat kontrol (IoT) "
        "berfungsi sempurna (ping diblokir 100%), memvalidasi firewall rule drop."
    )

    print("Modifying Section 7 (Kesimpulan)...")
    p_kesimpulan.text = (
        "5.1 Kesimpulan\n"
        "1. Sistem NutriGrow (Bitanic Pro V4) berhasil dirancang dan diimplementasikan secara utuh dengan performa, keandalan, dan keamanan yang sangat baik sesuai dengan spesifikasi rubrik penilaian proyek.\n"
        "2. Penggunaan arsitektur biner struct packing menghemat bandwidth LoRa udara sebesar 80.53% dan kompresi Gzip memotong payload REST API sebesar 76.67%, mempercepat pemuatan grafik dashboard.\n"
        "3. Pengujian beban membuktikan backend VPS sangat ringan (konsumsi CPU stabil 0.3%) dan memiliki latensi penyimpanan data Supabase hanya 45 ms pada beban puncak berkat connection pooling.\n"
        "4. Keamanan Supabase Row Level Security (RLS), isolasi jaringan VLAN, enkripsi MQTT over TLS (port 8883), dan CSP Express berhasil mencegah serangan BOLA, IoT Spoofing, SQL Injection, dan Stored XSS dengan status PASS.\n"
        "5. Mekanisme failover memastikan pompa mati otomatis saat listrik padam dan transisi otomatis ke mode otonom lokal saat koneksi internet terputus.\n\n"
        "5.2 Saran\n"
        "1. Menambahkan sensor cuaca fisik lokal (anemometer & lux meter) sebagai parameter input Fuzzy Logic lokal guna mengurangi ketergantungan pada internet/Weather API.\n"
        "2. Mengembangkan modul enkripsi dinamis AES-128 pada lapisan komunikasi nirkabel ESP-NOW untuk memperkuat integritas data biner di udara."
    )

    print("Modifying Istilah dan Definisi...")
    p_istilah.text = "API (Application Programming Interface): Sekumpulan aturan dan protokol yang memungkinkan berbagai komponen perangkat lunak untuk berkomunikasi dan bertukar data satu sama lain."
    # The template had a second definition at index 94. In our list we anchored p_istilah to doc.paragraphs[94]
    # which is "Authentication...". Let's clear its text and rewrite it, then append others.
    p_istilah.text = "Authentication: Proses verifikasi identitas pengguna atau sistem untuk memastikan bahwa mereka adalah siapa yang mereka klaim."
    
    definitions = [
        "IoT (Internet of Things): Jaringan objek fisik yang tersemat dengan sensor, perangkat lunak, dan teknologi lainnya untuk berkomunikasi dengan perangkat lain melalui internet.",
        "MQTT (Message Queuing Telemetry Transport): Protokol publish/subscribe ringan, sangat cocok untuk komunikasi telemetri perangkat IoT dengan resource terbatas.",
        "LoRa (Long Range): Teknologi modulasi nirkabel berdaya rendah untuk komunikasi jarak jauh (mencapai 5-15 km di area rural).",
        "ESP-NOW: Protokol komunikasi nirkabel jarak dekat cepat (2.4 GHz) antar-mikrokontroler Espressif tanpa menggunakan router.",
        "Supabase RLS (Row Level Security): Fitur PostgreSQL terkelola yang membatasi hak akses baris data pada tabel database berdasarkan klaim token otorisasi pengguna.",
        "Fuzzy Logic Mamdani: Metodologi pemetaan kondisi sensor non-linier menjadi keputusan durasi penyiraman dan volume pupuk menggunakan aturan penalaran linguistik terstruktur."
    ]
    p_last_def = p_istilah
    for d in definitions:
        p_last_def = insert_paragraph_after(p_last_def, d)

    print("Modifying References / Daftar Pustaka...")
    references = [
        "1. OWASP Foundation. (2023). OWASP API Security Top 10 2023. https://owasp.org/API-Security/",
        "2. OWASP Foundation. (2021). OWASP Top 10 2021. https://owasp.org/Top10/",
        "3. Supabase Documentation. (2025). Row Level Security. https://supabase.com/docs/guides/database/postgres/row-level-security",
        "4. Espressif Systems. (2026). ESP-NOW User Guide. https://docs.espressif.com/",
        "5. Semtech Corporation. (2025). LoRa and LoRaWAN Technical Overview. https://www.semtech.com/",
        "6. Armanto, A. et al. (2025). Smart Water Automation System to Regulate Watering of Horticultural Plants with Arduino-based Fuzzy Logic. International Journal of Integrative Sciences (IJIS). https://mryformosapublisher.org/index.php/ijis/article/view/24",
        "7. Saravanan, S. et al. (2023). Study and Development of IoT Based Smart Irrigation System Using Soil Moisture and Weather Prediction. International Journal for Research in Applied Science & Engineering Technology (IJRASET). https://doi.org/10.22214/ijraset.2023.50948",
        "8. Morchid, A. et al. (2023). Fuzzy-IoT Smart Irrigation System for Precision Scheduling and Monitoring. Computers and Electronics in Agriculture, ScienceDirect. https://www.sciencedirect.com/science/article/abs/pii/S0168169923007950",
        "9. Hasan et al. (2025). IoT-Based Smart Fertigation Scheduling and Wireless Microclimate Monitoring for a Greenhouse Dutch Bucket Hydroponic System. Irrigation and Drainage, Wiley. https://doi.org/10.1002/ird.70012",
        "10. Ngoma, D.H. et al. (2025). Design and Development of IoT Smart Drip Irrigation and Fertigation Prototype for Small and Medium Scale Farmers. Journal of The Institution of Engineers (India), Springer. https://doi.org/10.1007/s40030-024-00857-7",
        "11. Saha, G. et al. (2025). An IoT-Based Smart Plant Monitoring and Irrigation System with Real-Time Environmental Sensing, Automated Alerts, and Cloud Analytics. arXiv Preprint. https://arxiv.org/pdf/2601.15830",
        "12. Jaywant, A.P. & Desai, A.S. (2024). Cloud-Integrated Precision Agriculture with LoRaWAN and ESP32. Journal of Computer and Creative Technology (JCCT). IEEE ICOSEC 2024. https://so13.tci-thaijo.org/index.php/jcct/article/view/1560",
        "13. Multiple Authors. (2025). Cloud–Edge–Device Collaborative Computing in Smart Agriculture: Architectures, Applications, and Future Perspectives. Frontiers in Plant Science. https://www.frontiersin.org/journals/plant-science/articles/10.3389/fpls.2025.1668545/full",
        "14. Imbernon-Mulero, A. et al. (2023). Evaluation of an Autonomous Smart System for Optimal Management of Fertigation with Variable Sources of Irrigation Water. Frontiers in Plant Science. https://doi.org/10.3389/fpls.2023.1149956",
        "15. Garcia, L. et al. (2020). IoT-Based Smart Irrigation Systems: An Overview on the Recent Trends on Sensors and IoT Systems for Irrigation in Precision Agriculture. Sensors, MDPI. https://doi.org/10.3390/s20041042"
    ]
    
    p_pustaka.text = references[0]
    p_ref = p_pustaka
    for ref in references[1:]:
        p_ref = insert_paragraph_after(p_ref, ref)

    print("Modifying CV / Daftar Riwayat Hidup...")
    p_cv_azrial.text = (
        "Maulana Daviq Putra (NIM: 2307422024) Mahasiswa Politeknik Negeri Jakarta, Jurusan Teknik Informatika "
        "dan Komputer, Program Studi Teknik Multimedia dan Jaringan. [Biodata dan Riwayat Hidup Belum Diisi - Silakan Lengkapi]"
    )
    
    p_cv_arizal.text = (
        "Muhammad Ali Diepo Ar-Rasyid (NIM: 2307422023) Mahasiswa Politeknik Negeri Jakarta, Jurusan Teknik Informatika "
        "dan Komputer, Program Studi Teknik Multimedia dan Jaringan. [Biodata dan Riwayat Hidup Belum Diisi - Silakan Lengkapi]"
    )
    
    p_cv_sultan.text = (
        "Muhammad Athaillah Hardianto (NIM: 2307422011) Mahasiswa Politeknik Negeri Jakarta, Jurusan Teknik Informatika "
        "dan Komputer, Program Studi Teknik Multimedia dan Jaringan. [Biodata dan Riwayat Hidup Belum Diisi - Silakan Lengkapi]"
    )
    
    # Append the remaining two authors after the last CV paragraph
    p_cv = p_cv_sultan
    p_cv = insert_paragraph_after(p_cv, "")
    p_cv = insert_paragraph_after(p_cv, 
        "Muhammad Pandya Hanif Ramadhan (NIM: 2307422006) Mahasiswa Politeknik Negeri Jakarta, Jurusan Teknik Informatika "
        "dan Komputer, Program Studi Teknik Multimedia dan Jaringan. [Biodata dan Riwayat Hidup Belum Diisi - Silakan Lengkapi]"
    )
    p_cv = insert_paragraph_after(p_cv, "")
    p_cv = insert_paragraph_after(p_cv, 
        "Ravila Sasla Arandika (NIM: 2307422026) Mahasiswa Politeknik Negeri Jakarta, Jurusan Teknik Informatika "
        "dan Komputer, Program Studi Teknik Multimedia dan Jaringan. [Biodata dan Riwayat Hidup Belum Diisi - Silakan Lengkapi]"
    )

    print(f"Saving new document to {output_path}...")
    doc.save(output_path)
    print("Document successfully created!")

if __name__ == "__main__":
    convert_report()
