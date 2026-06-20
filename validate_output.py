import docx
import os

def validate():
    path = "Laporan_Akhir_NutriGrow_Baru_v2.docx"
    if not os.path.exists(path):
        print(f"Error: {path} not found.")
        return
        
    print(f"Opening {path} for validation...")
    doc = docx.Document(path)
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    
    # Check cover page
    print("\nCover Page Paragraphs:")
    for i in range(12):
        if i < len(doc.paragraphs):
            print(f"  P{i}: {repr(doc.paragraphs[i].text)}")
            
    # Check some headings and contents
    print("\nSample Heading check:")
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip():
            print(f"  P{idx} ({p.style.name}): {p.text.strip()}")
            
    # Check tables
    print("\nTables check:")
    for idx, t in enumerate(doc.tables):
        print(f"  Table {idx+1}: {len(t.rows)} rows, {len(t.columns)} cols")
        # print first row
        first_row_cells = [cell.text.strip().replace('\n', ' ') for cell in t.rows[0].cells]
        print(f"    Header: {first_row_cells}")
        
    print("\nValidation completed successfully!")

if __name__ == "__main__":
    validate()
